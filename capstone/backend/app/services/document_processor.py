"""
Document Processing Service for Legal Documents

This service handles parsing of legal documents with positional information
for precise clause identification and amendment tracking.
"""

import io
import json
import logging
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass
from pathlib import Path

# PDF Processing
import PyPDF2
from pdf2image import convert_from_bytes
import pytesseract
from PIL import Image

# DOCX Processing  
from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run

# Async support
import asyncio
from concurrent.futures import ThreadPoolExecutor

logger = logging.getLogger(__name__)


@dataclass
class DocumentPosition:
    """Represents a position within a document"""
    page_number: int
    paragraph_index: int
    char_start: int
    char_end: int
    bbox: Optional[Dict[str, float]] = None  # Bounding box for PDF


@dataclass
class DocumentChunk:
    """Represents a chunk of text with positional information"""
    text: str
    position: DocumentPosition
    chunk_type: str  # 'paragraph', 'heading', 'clause', etc.
    metadata: Dict[str, Any]


@dataclass
class ProcessedDocument:
    """Represents a fully processed document"""
    document_id: str
    title: str
    content: str
    chunks: List[DocumentChunk]
    metadata: Dict[str, Any]
    processing_stats: Dict[str, Any]


class DocumentProcessor:
    """Main document processor for legal documents"""
    
    def __init__(self):
        self.executor = ThreadPoolExecutor(max_workers=2)
        
    async def process_document(
        self,
        file_content: bytes,
        filename: str,
        document_id: str,
        file_type: str
    ) -> ProcessedDocument:
        """
        Process a document and extract structured content with positions
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            document_id: Unique document identifier
            file_type: File extension (pdf, docx, txt)
            
        Returns:
            ProcessedDocument with structured content and positions
        """
        try:
            logger.info(f"Processing document {document_id} of type {file_type}")
            
            if file_type.lower() == 'pdf':
                return await self._process_pdf(file_content, filename, document_id)
            elif file_type.lower() == 'docx':
                return await self._process_docx(file_content, filename, document_id)
            elif file_type.lower() == 'txt':
                return await self._process_txt(file_content, filename, document_id)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            logger.error(f"Error processing document {document_id}: {str(e)}")
            raise
    
    async def _process_pdf(
        self,
        file_content: bytes,
        filename: str,
        document_id: str
    ) -> ProcessedDocument:
        """Process PDF with positional information"""
        
        def _extract_pdf_content():
            try:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                chunks = []
                full_text = ""
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    
                    if not page_text.strip():
                        # Fallback to OCR for scanned PDFs
                        page_text = self._ocr_pdf_page(file_content, page_num)
                    
                    # Split into paragraphs
                    paragraphs = self._split_into_paragraphs(page_text)
                    
                    char_offset = len(full_text)
                    
                    for para_idx, paragraph in enumerate(paragraphs):
                        if paragraph.strip():
                            position = DocumentPosition(
                                page_number=page_num + 1,
                                paragraph_index=para_idx,
                                char_start=char_offset,
                                char_end=char_offset + len(paragraph)
                            )
                            
                            chunk_type = self._identify_chunk_type(paragraph)
                            
                            chunk = DocumentChunk(
                                text=paragraph,
                                position=position,
                                chunk_type=chunk_type,
                                metadata={
                                    "page_number": page_num + 1,
                                    "paragraph_index": para_idx,
                                    "word_count": len(paragraph.split()),
                                    "char_count": len(paragraph)
                                }
                            )
                            chunks.append(chunk)
                            
                            full_text += paragraph + "\n"
                            char_offset = len(full_text)
                
                return full_text, chunks, {
                    "total_pages": len(pdf_reader.pages),
                    "extraction_method": "text_extraction"
                }
                
            except Exception as e:
                logger.error(f"PDF processing error: {str(e)}")
                raise
        
        # Run in thread pool for CPU-intensive operations
        full_text, chunks, stats = await asyncio.get_event_loop().run_in_executor(
            self.executor, _extract_pdf_content
        )
        
        return ProcessedDocument(
            document_id=document_id,
            title=filename,
            content=full_text,
            chunks=chunks,
            metadata={
                "file_type": "pdf",
                "filename": filename,
                "total_chunks": len(chunks)
            },
            processing_stats=stats
        )
    
    async def _process_docx(
        self,
        file_content: bytes,
        filename: str,
        document_id: str
    ) -> ProcessedDocument:
        """Process DOCX with structure preservation"""
        
        def _extract_docx_content():
            try:
                doc = Document(io.BytesIO(file_content))
                chunks = []
                full_text = ""
                char_offset = 0
                
                for para_idx, paragraph in enumerate(doc.paragraphs):
                    if paragraph.text.strip():
                        position = DocumentPosition(
                            page_number=1,  # DOCX doesn't have clear page breaks
                            paragraph_index=para_idx,
                            char_start=char_offset,
                            char_end=char_offset + len(paragraph.text)
                        )
                        
                        chunk_type = self._identify_chunk_type_docx(paragraph)
                        
                        # Extract formatting metadata
                        formatting_info = self._extract_docx_formatting(paragraph)
                        
                        chunk = DocumentChunk(
                            text=paragraph.text,
                            position=position,
                            chunk_type=chunk_type,
                            metadata={
                                "paragraph_index": para_idx,
                                "word_count": len(paragraph.text.split()),
                                "char_count": len(paragraph.text),
                                "formatting": formatting_info
                            }
                        )
                        chunks.append(chunk)
                        
                        full_text += paragraph.text + "\n"
                        char_offset = len(full_text)
                
                return full_text, chunks, {
                    "total_paragraphs": len(doc.paragraphs),
                    "extraction_method": "docx_structure"
                }
                
            except Exception as e:
                logger.error(f"DOCX processing error: {str(e)}")
                raise
        
        full_text, chunks, stats = await asyncio.get_event_loop().run_in_executor(
            self.executor, _extract_docx_content
        )
        
        return ProcessedDocument(
            document_id=document_id,
            title=filename,
            content=full_text,
            chunks=chunks,
            metadata={
                "file_type": "docx",
                "filename": filename,
                "total_chunks": len(chunks)
            },
            processing_stats=stats
        )
    
    async def _process_txt(
        self,
        file_content: bytes,
        filename: str,
        document_id: str
    ) -> ProcessedDocument:
        """Process plain text files"""
        
        try:
            # Decode text content
            content = file_content.decode('utf-8')
            
            # Split into paragraphs
            paragraphs = self._split_into_paragraphs(content)
            chunks = []
            char_offset = 0
            
            for para_idx, paragraph in enumerate(paragraphs):
                if paragraph.strip():
                    position = DocumentPosition(
                        page_number=1,
                        paragraph_index=para_idx,
                        char_start=char_offset,
                        char_end=char_offset + len(paragraph)
                    )
                    
                    chunk_type = self._identify_chunk_type(paragraph)
                    
                    chunk = DocumentChunk(
                        text=paragraph,
                        position=position,
                        chunk_type=chunk_type,
                        metadata={
                            "paragraph_index": para_idx,
                            "word_count": len(paragraph.split()),
                            "char_count": len(paragraph)
                        }
                    )
                    chunks.append(chunk)
                    char_offset += len(paragraph) + 1  # +1 for newline
            
            return ProcessedDocument(
                document_id=document_id,
                title=filename,
                content=content,
                chunks=chunks,
                metadata={
                    "file_type": "txt",
                    "filename": filename,
                    "total_chunks": len(chunks)
                },
                processing_stats={
                    "total_paragraphs": len(paragraphs),
                    "extraction_method": "text_decode"
                }
            )
            
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    content = file_content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise ValueError("Unable to decode text file")
    
    def _ocr_pdf_page(self, file_content: bytes, page_num: int) -> str:
        """Perform OCR on a PDF page for scanned documents"""
        try:
            # Convert PDF page to image
            images = convert_from_bytes(
                file_content,
                first_page=page_num + 1,
                last_page=page_num + 1,
                dpi=200
            )
            
            if images:
                # Perform OCR
                text = pytesseract.image_to_string(images[0])
                return text
            
            return ""
            
        except Exception as e:
            logger.warning(f"OCR failed for page {page_num}: {str(e)}")
            return ""
    
    def _split_into_paragraphs(self, text: str) -> List[str]:
        """Split text into meaningful paragraphs"""
        # Split on double newlines first
        paragraphs = text.split('\n\n')
        
        # Further split if paragraphs are too long
        final_paragraphs = []
        for para in paragraphs:
            if len(para) > 1000:  # Split long paragraphs
                sentences = para.split('. ')
                current_para = ""
                for sentence in sentences:
                    if len(current_para + sentence) > 1000 and current_para:
                        final_paragraphs.append(current_para.strip())
                        current_para = sentence + ". "
                    else:
                        current_para += sentence + ". "
                if current_para:
                    final_paragraphs.append(current_para.strip())
            else:
                final_paragraphs.append(para.strip())
        
        return [p for p in final_paragraphs if p]
    
    def _identify_chunk_type(self, text: str) -> str:
        """Identify the type of text chunk for legal documents"""
        text_lower = text.lower().strip()
        
        # Legal document patterns
        if any(word in text_lower for word in ['article', 'section', 'clause', 'paragraph']):
            if text_lower.startswith(('article', 'section', 'clause')):
                return 'heading'
            return 'clause'
        
        # Check for definitions
        if 'definitions' in text_lower or text_lower.startswith('for purposes of'):
            return 'definition'
        
        # Check for numbered/lettered lists
        if text_lower.startswith(('(a)', '(b)', '(1)', '(2)', 'a.', 'b.', '1.', '2.')):
            return 'list_item'
        
        # Default to paragraph
        return 'paragraph'
    
    def _identify_chunk_type_docx(self, paragraph: Paragraph) -> str:
        """Identify chunk type for DOCX paragraphs using formatting"""
        # Check heading styles
        if paragraph.style.name.startswith('Heading'):
            return 'heading'
        
        # Check for bold text (potential headings)
        if paragraph.runs and any(run.bold for run in paragraph.runs):
            return 'heading'
        
        # Use text-based identification
        return self._identify_chunk_type(paragraph.text)
    
    def _extract_docx_formatting(self, paragraph: Paragraph) -> Dict[str, Any]:
        """Extract formatting information from DOCX paragraph"""
        formatting = {
            "style": paragraph.style.name,
            "bold": False,
            "italic": False,
            "underline": False
        }
        
        if paragraph.runs:
            formatting["bold"] = any(run.bold for run in paragraph.runs)
            formatting["italic"] = any(run.italic for run in paragraph.runs)
            formatting["underline"] = any(run.underline for run in paragraph.runs)
        
        return formatting
